import docx
import io
import base64
from typing import Dict, Any, List
from .text_normalizer import normalize_text

def process_docx(file_path: str, extract_images: bool = True) -> Dict[str, Any]:
    """
    Processes DOCX files using python-docx:
    1. Extracts headings, paragraphs, and tables as formatted markdown text.
    2. Extracts embedded images if visual understanding is needed.
    Returns:
    {
      "text": str,
      "embedded_images": [{"id": str, "type": "embedded_image", "base64_image": str}],
      "metadata": {"paragraphs": int, "tables": int, "images_count": int}
    }
    """
    doc = docx.Document(file_path)
    text_blocks = []
    
    # 1. Paragraphs & Headings
    para_count = 0
    for p in doc.paragraphs:
        p_text = p.text.strip()
        if not p_text:
            continue
        para_count += 1
        style_name = p.style.name.lower() if p.style else ""
        if "heading 1" in style_name:
            text_blocks.append(f"# {p_text}")
        elif "heading 2" in style_name:
            text_blocks.append(f"## {p_text}")
        elif "heading 3" in style_name:
            text_blocks.append(f"### {p_text}")
        else:
            text_blocks.append(p_text)
            
    # 2. Tables
    table_count = 0
    for table in doc.tables:
        table_count += 1
        rows_str = []
        for i, row in enumerate(table.rows):
            cell_texts = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            rows_str.append(" | ".join(cell_texts))
            if i == 0:
                rows_str.append(" | ".join(["---"] * len(cell_texts)))
        if rows_str:
            text_blocks.append(f"\n### Table {table_count}\n" + "\n".join(rows_str))

    # 3. Embedded Images
    embedded_images = []
    if extract_images:
        try:
            for rel_id, rel in doc.part.rels.items():
                if "image" in rel.target_ref:
                    img_part = rel.target_part
                    img_bytes = img_part.blob
                    b64_img = base64.b64encode(img_bytes).decode("utf-8")
                    embedded_images.append({
                        "id": rel_id,
                        "type": "embedded_image",
                        "base64_image": b64_img
                    })
        except Exception as img_err:
            print(f"[DOCX Processor] Image extraction warning: {img_err}")

    full_text = normalize_text("\n\n".join(text_blocks))
    
    return {
        "text": full_text,
        "embedded_images": embedded_images,
        "metadata": {
            "paragraphs": para_count,
            "tables": table_count,
            "images_count": len(embedded_images)
        }
    }
